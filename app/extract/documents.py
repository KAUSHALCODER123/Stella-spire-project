"""Turn an uploaded CV into plain text.

Deliberately dumb. This layer does no interpretation at all -- it only gets
characters out of the container, records what it could not do, and hands the
result to the model. Every heuristic that could silently mangle a CV (column
reordering, header stripping, bullet normalising) is left out on purpose,
because a wrong guess here is invisible downstream.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}

# Below this many characters per page, the page is almost certainly a scan or
# an image-only export and the model will be reading nothing.
MIN_CHARS_PER_PAGE = 120


@dataclass
class DocumentText:
    text: str
    page_count: int
    source_format: str
    warnings: List[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.text)


def _clean(raw: str) -> str:
    # Collapse the runs of blank lines that PDF extraction leaves behind,
    # without touching intra-line spacing (which can carry column structure).
    raw = raw.replace("\x00", "")
    raw = re.sub(r"[ \t]+\n", "\n", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()


def _from_pdf(path: Path) -> DocumentText:
    import fitz  # PyMuPDF

    warnings: List[str] = []
    parts: List[str] = []
    with fitz.open(path) as doc:
        page_count = doc.page_count
        for page in doc:
            # "text" mode applies PyMuPDF's reading-order heuristic. It is good
            # on single-column CVs and unreliable on two-column templates,
            # which is a known and measured weakness -- see eval/README.md.
            parts.append(page.get_text("text"))

    text = _clean("\n\n".join(parts))

    if page_count and len(text) / page_count < MIN_CHARS_PER_PAGE:
        warnings.append(
            "Very little extractable text ({} chars across {} pages). This is most likely a scanned "
            "or image-based PDF; the dossier will be unreliable. Ask the candidate for a text CV.".format(
                len(text), page_count
            )
        )

    return DocumentText(text=text, page_count=page_count, source_format="pdf", warnings=warnings)


def _from_docx(path: Path) -> DocumentText:
    import docx

    document = docx.Document(str(path))
    parts: List[str] = [p.text for p in document.paragraphs]

    # Plenty of CV templates hold the entire work history inside a table, and
    # python-docx does not surface those through .paragraphs.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return DocumentText(text=_clean("\n".join(parts)), page_count=0, source_format="docx")


def _from_text(path: Path) -> DocumentText:
    raw = path.read_text(encoding="utf-8", errors="replace")
    return DocumentText(text=_clean(raw), page_count=0, source_format=path.suffix.lstrip("."))


def extract_text(path: str | Path) -> DocumentText:
    """Read a CV file into text. Raises ValueError for unsupported types."""
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(
            "Unsupported file type '{}'. Supported: {}.".format(suffix or "(none)", ", ".join(sorted(SUPPORTED_SUFFIXES)))
        )
    if not path.exists():
        raise FileNotFoundError(str(path))

    if suffix == ".pdf":
        result = _from_pdf(path)
    elif suffix == ".docx":
        result = _from_docx(path)
    else:
        result = _from_text(path)

    if not result.text.strip():
        result.warnings.append("No text could be read from this document at all.")

    return result
